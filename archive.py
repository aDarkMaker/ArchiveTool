import os
import re
import sys
import requests
import tkinter as tk
from datetime import datetime
from threading import Thread, Event
from tkinter import ttk, scrolledtext, filedialog
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from urllib.parse import urlparse
import warnings
from webdriver_manager.chrome import ChromeDriverManager

warnings.filterwarnings("ignore", category=DeprecationWarning)

# ------------------------- 核心功能函数 -------------------------
def sanitize_filename(filename):
    """移除文件名中的非法字符"""
    return re.sub(r'[\\/:*?"<>|]', '', filename)

def get_web_content(url):
    """使用Selenium获取动态网页内容"""
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
    chrome_options.add_argument("accept-language=zh-CN,zh;q=0.9")
    
    try:
        driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()),
            options=chrome_options
        )
        driver.implicitly_wait(15)
        
        driver.get(url)
        WebDriverWait(driver, 15).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, '#publish_time, .rich_media_meta_text'))
        )
        return driver.page_source
    except Exception as e:
        print(f"页面加载异常: {str(e)}")
        return None
    finally:
        if 'driver' in locals():
            driver.quit()

def parse_date(publish_time_str):
    """增强的日期解析函数"""
    patterns = [
        r'(\d{4})年(\d{1,2})月(\d{1,2})日',
        r'(\d{4})-(\d{2})-(\d{2})',
        r'(\d{4})/(\d{2})/(\d{2})'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, publish_time_str)
        if match:
            try:
                year = int(match.group(1))
                month = int(match.group(2))
                day = int(match.group(3))
                return f"{year:04}{month:02}{day:02}"
            except:
                continue
                
    print(f"无法解析日期，使用当前日期: {publish_time_str}")
    return datetime.now().strftime("%Y%m%d")

def parse_wechat_article(html):
    """解析微信公众号文章"""
    soup = BeautifulSoup(html, 'lxml')
    
    publish_time_tag = soup.find('em', {'id': 'publish_time'}) or \
                      soup.find('em', class_=re.compile('rich_media_meta')) or \
                      soup.find('div', class_=re.compile('rich_media_meta_text'))
    
    date_str = parse_date(publish_time_tag.get_text()) if publish_time_tag else datetime.now().strftime("%Y%m%d")
    
    title = "无标题"
    og_title = soup.find('meta', {'property': 'og:title'})
    if og_title and og_title.get('content'):
        title = og_title['content'].strip()
    else:
        title_selector = soup.select_one('h1.rich_media_title, #activity-name, title')
        if title_selector:
            title = title_selector.get_text().strip()
    
    content_div = soup.select_one('div.rich_media_content, #js_content') 
    
    content_paragraphs = []
    imgs = []
    
    if content_div:
        elements = content_div.select('p, img, span, strong, section')
        for element in elements:
            if element.name == 'p':
                content_paragraphs.append(element)
            elif element.name == 'img':
                img_url = element.get('data-src') or element.get('src')
                if img_url and not img_url.startswith('data:'):
                    imgs.append(img_url)
    
    return date_str, title, content_paragraphs, imgs

def set_doc_style(doc):
    """设置文档样式"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_formatted_paragraph(doc, paragraph):
    """添加格式化段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    
    try:
        for content in paragraph.contents:
            if isinstance(content, str):
                run = p.add_run(content.strip())
            elif content.name == 'strong':
                run = p.add_run(content.get_text().strip())
                run.bold = True
            elif content.name == 'span':
                run = p.add_run(content.get_text().strip())
                if 'color:' in content.get('style', ''):
                    color_str = re.findall(r'color:(#[0-9a-fA-F]+)', content['style'])
                    if color_str:
                        rgb = tuple(int(color_str[0][i:i+2], 16) for i in (1, 3, 5))
                        run.font.color.rgb = RGBColor(*rgb)
        
        if not p.text.strip():
            p._element.getparent().remove(p._element)
    except Exception as e:
        p._element.getparent().remove(p._element)

def remove_audit_info(doc):
    """删除审核信息段落"""
    audit_pattern = re.compile(r'审核\s*\|\s*.+')
    for para in reversed(doc.paragraphs):
        if audit_pattern.match(para.text.strip()):
            p = para._element
            p.getparent().remove(p)
            p._p = p._element = None

def download_image(img_url, save_path):
    """下载图片"""
    headers = {
        'User-Agent': 'Mozilla/5.0',
        'Referer': 'https://mp.weixin.qq.com/'
    }
    for retry in range(3):
        try:
            response = requests.get(img_url, headers=headers, stream=True, timeout=20)
            if response.status_code == 200:
                with open(save_path, 'wb') as f:
                    for chunk in response.iter_content(1024):
                        f.write(chunk)
                return True
        except Exception as e:
            if retry == 2:
                print(f"图片下载失败: {img_url}")
                return False
            print(f"重试下载({retry+1}/3): {img_url}")
    return False

# ------------------------- GUI界面 -------------------------
class WeChatDownloaderGUI:
    def __init__(self, master):
        self.master = master
        master.title("推文归档Helper 测试版")
        master.geometry("800x700")
        
        # 图标加载优化
        self.load_application_icon()
        
        master.resizable(False, False)
        self.stop_event = Event()
        self.create_widgets()
        self.original_stdout = sys.stdout
        sys.stdout = self
        
        # 窗口关闭事件绑定
        master.protocol("WM_DELETE_WINDOW", self.on_close)

    def load_application_icon(self):
        """加载应用程序图标"""
        try:
            if getattr(sys, 'frozen', False):
                # 打包后的路径
                base_path = sys._MEIPASS
            else:
                # 开发环境路径
                base_path = os.path.dirname(os.path.abspath(__file__))
            
            icon_path = os.path.join(base_path, 'icon.ico')
            print(f"尝试加载图标路径: {icon_path}")
            
            if os.path.exists(icon_path):
                self.master.iconbitmap(icon_path)
            else:
                print("警告: 图标文件未找到，使用默认图标")
        except Exception as e:
            print(f"图标加载错误: {str(e)}")

    def on_close(self):
        """窗口关闭时的清理操作"""
        self.stop_event.set()
        self.master.destroy()
        sys.stdout = self.original_stdout
        os._exit(0)

    def create_widgets(self):
        # URL输入
        url_frame = ttk.Frame(self.master)
        url_frame.pack(pady=10, padx=10, fill=tk.X)
        ttk.Label(url_frame, text="文章URL:").pack(side=tk.LEFT)
        self.url_entry = ttk.Entry(url_frame, width=60)
        self.url_entry.pack(side=tk.LEFT, padx=5)
        
        # 路径选择
        path_frame = ttk.Frame(self.master)
        path_frame.pack(pady=5, padx=10, fill=tk.X)
        ttk.Label(path_frame, text="保存路径:").pack(side=tk.LEFT)
        self.path_var = tk.StringVar()
        ttk.Entry(path_frame, textvariable=self.path_var, width=50).pack(side=tk.LEFT, padx=5)
        ttk.Button(path_frame, text="浏览...", command=self.select_path).pack(side=tk.LEFT)
        
        # 日志区域
        log_frame = ttk.Frame(self.master)
        log_frame.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)
        self.log_area = scrolledtext.ScrolledText(
            log_frame, wrap=tk.WORD, font=('微软雅黑', 10), bg='#333', fg='#fff')
        self.log_area.pack(fill=tk.BOTH, expand=True)
        
        # 控制按钮
        btn_frame = ttk.Frame(self.master)
        btn_frame.pack(pady=10)
        self.start_btn = ttk.Button(btn_frame, text="开始下载", command=self.start_download)
        self.start_btn.pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="清空日志", command=self.clear_log).pack(side=tk.LEFT, padx=5)

    def select_path(self):
        path = filedialog.askdirectory()
        if path:
            self.path_var.set(path)

    def write(self, text):
        self.log_area.insert(tk.END, text)
        self.log_area.see(tk.END)

    def flush(self):
        pass

    def clear_log(self):
        self.log_area.delete(1.0, tk.END)

    def start_download(self):
        url = self.url_entry.get().strip()
        base_dir = self.path_var.get().strip()
        
        if not re.match(r'^https?://mp\.weixin\.qq\.com/s\S*', url):
            self.log_area.insert(tk.END, "错误：请输入有效的微信公众号文章链接\n")
            return
        
        if not base_dir:
            base_dir = filedialog.askdirectory()
            if not base_dir:
                return
            self.path_var.set(base_dir)
        
        self.start_btn.config(state=tk.DISABLED)
        
        def run_download():
            try:
                modified_main(url, base_dir)
            except Exception as e:
                print(f"发生错误: {str(e)}")
            finally:
                self.start_btn.config(state=tk.NORMAL)
        
        Thread(target=run_download, daemon=True).start()

    def __del__(self):
        sys.stdout = self.original_stdout

def modified_main(url, base_dir):
    """带中断检查的主逻辑"""
    try:
        print("="*50 + "\n开始处理文章...")
        html = get_web_content(url)
        if not html:
            print("错误：无法获取网页内容")
            return
        
        date_str, title, content_paragraphs, imgs = parse_wechat_article(html)
        
        print(f"解析结果：\n日期：{date_str}\n标题：{title[:50]}...\n段落数：{len(content_paragraphs)}\n图片数：{len(imgs)}\n")
        
        folder_name = f"{date_str}_{sanitize_filename(title)}"
        folder_path = os.path.join(base_dir, folder_name)
        img_dir = os.path.join(folder_path, '图片')
        os.makedirs(img_dir, exist_ok=True)
        
        doc = Document()
        set_doc_style(doc)
        
        for p in content_paragraphs:
            try:
                add_formatted_paragraph(doc, p)
            except Exception as e:
                print(f"跳过无效段落: {str(e)}\n")
        
        remove_audit_info(doc)
        
        doc_path = os.path.join(folder_path, '文字.docx')
        doc.save(doc_path)
        print(f"文档保存成功：{doc_path}\n")
        
        success_count = 0
        for idx, img_url in enumerate(imgs, 1):
            ext = os.path.splitext(urlparse(img_url).path)[1] or '.jpg'
            save_path = os.path.join(img_dir, f'图片{idx}{ext}')
            if download_image(img_url, save_path):
                success_count += 1
            print(f"图片下载进度：{idx}/{len(imgs)}")
        
        print(f"图片下载完成：成功 {success_count}/{len(imgs)}")
        print(f"处理完成！保存路径：{os.path.abspath(folder_path)}\n" + "="*50 + "\n")
    
    except Exception as e:
        print(f"严重错误: {str(e)}\n")

if __name__ == "__main__":
    root = tk.Tk()
    app = WeChatDownloaderGUI(root)
    root.mainloop()